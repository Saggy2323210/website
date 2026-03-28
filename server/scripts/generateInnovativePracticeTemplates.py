#!/usr/bin/env python3
"""
Generate DOCX templates for Innovative Practice markdown conversion
Templates show the expected markdown table format for admins to edit in Word
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Define output directory
OUTPUT_DIR = os.path.join(
    os.path.dirname(__file__),
    "..",
    "uploads",
    "documents",
    "innovative_practice_templates",
)

# Create directory if it doesn't exist
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Department data
DEPARTMENTS = {
    "CSE": {
        "name": "Computer Science and Engineering",
        "sample_rows": [
            ("01", "Dr. Faculty Name", "Subject Name", "Practice Method", "/uploads/documents/path.pdf"),
            ("02", "Prof. Another Name", "Another Subject", "Another Method", "https://example.com"),
        ]
    },
    "IT": {
        "name": "Information Technology",
        "sample_rows": [
            ("01", "Dr. Faculty Name", "Subject Name", "Practice Method", "/uploads/documents/path.pdf"),
            ("02", "Prof. Another Name", "Another Subject", "Another Method", "https://example.com"),
        ]
    },
    "Mechanical": {
        "name": "Mechanical Engineering",
        "sample_rows": [
            ("01", "Dr. Faculty Name", "Subject Name", "Practice Method", "/uploads/documents/path.pdf"),
            ("02", "Prof. Another Name", "Another Subject", "Another Method", "https://example.com"),
        ]
    },
    "EnTC": {
        "name": "Electronics and Telecommunication Engineering",
        "sample_rows": [
            ("01", "Dr. Faculty Name", "Subject Name", "Practice Method", "/uploads/documents/path.pdf"),
            ("02", "Prof. Another Name", "Another Subject", "Another Method", "https://example.com"),
        ]
    },
    "Electrical": {
        "name": "Electrical Engineering",
        "sample_rows": [
            ("01", "Dr. Faculty Name", "Subject Name", "Practice Method", "/uploads/documents/path.pdf"),
        ]
    },
    "MBA": {
        "name": "Master of Business Administration",
        "sample_rows": [
            ("01", "Dr. Faculty Name", "Subject Name", "Practice Method", "/uploads/documents/path.pdf"),
        ]
    },
}

def create_template(dept_key, dept_info):
    """Create a DOCX template for a department's innovative practices"""
    doc = Document()
    
    # Add title
    title = doc.add_heading(
        f"Innovative Practices in Teaching and Learning - {dept_info['name']}",
        level=1,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add instructions
    doc.add_heading("Instructions:", level=2)
    instructions = doc.add_paragraph()
    instructions.add_run("1. Edit the table below to add or update innovative practices\n")
    instructions.add_run("2. Keep the table structure intact with header row and separator\n")
    instructions.add_run("3. Add new rows following the same format\n")
    instructions.add_run("4. S.N. = Serial Number\n")
    instructions.add_run("5. Link = Internal path (/uploads/...) or external URL (https://...)\n")
    instructions.add_run("6. Save and import in the admin editor using DOCX import feature\n")
    
    doc.add_paragraph()
    
    # Add table
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    
    # Add header row
    header_cells = table.rows[0].cells
    headers = ["S.N.", "Faculty", "Subject", "Practice", "Link"]
    for i, header_text in enumerate(headers):
        header_cells[i].text = header_text
        # Make header bold
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Add sample rows
    for sn, faculty, subject, practice, link in dept_info["sample_rows"]:
        row_cells = table.add_row().cells
        row_cells[0].text = sn
        row_cells[1].text = faculty
        row_cells[2].text = subject
        row_cells[3].text = practice
        row_cells[4].text = link
    
    # Set column widths
    for row in table.rows:
        row.cells[0].width = Inches(0.8)
        row.cells[1].width = Inches(1.5)
        row.cells[2].width = Inches(2.0)
        row.cells[3].width = Inches(2.0)
        row.cells[4].width = Inches(2.0)
    
    # Save document
    dept_lower = dept_key.lower()
    filename = f"{dept_lower}_template.docx"
    filepath = os.path.join(OUTPUT_DIR, filename)
    doc.save(filepath)
    print(f"✓ Created: {filename}")

def main():
    """Generate all templates"""
    print("Generating Innovative Practice DOCX templates...")
    print(f"Output directory: {OUTPUT_DIR}\n")
    
    for dept_key, dept_info in DEPARTMENTS.items():
        create_template(dept_key, dept_info)
    
    print(f"\n✓ All templates created successfully in {OUTPUT_DIR}")

if __name__ == "__main__":
    main()
